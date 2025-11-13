#!/usr/bin/env python3
"""
Convert the PanSSR manuscript from Markdown to Word format with proper formatting.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def parse_markdown_file(md_file):
    """Parse the markdown file and return structured content."""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    return content

def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic)."""
    # Pattern to match **bold**, *italic*, or plain text
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`|[^*`]+)'
    parts = re.findall(pattern, text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Code text
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Plain text
            paragraph.add_run(part)

def create_word_document(md_content, output_file):
    """Create a formatted Word document from markdown content."""
    doc = Document()

    # Set up document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Split content into lines
    lines = md_content.split('\n')

    i = 0
    in_table = False
    table_lines = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'Normal'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    p.paragraph_format.left_indent = Inches(0.5)
                    code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Handle headings
        if line.startswith('# '):
            # Title (H1)
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(16)
                run.bold = True

        elif line.startswith('## '):
            # Section heading (H2)
            heading = line[3:].strip()
            p = doc.add_heading(heading, level=1)
            for run in p.runs:
                run.font.size = Pt(14)
                run.bold = True

        elif line.startswith('### '):
            # Subsection heading (H3)
            heading = line[4:].strip()
            p = doc.add_heading(heading, level=2)
            for run in p.runs:
                run.font.size = Pt(12)
                run.bold = True

        elif line.startswith('#### '):
            # Sub-subsection heading (H4)
            heading = line[5:].strip()
            p = doc.add_heading(heading, level=3)
            for run in p.runs:
                run.font.size = Pt(12)
                run.bold = True
                run.italic = True

        # Handle tables
        elif '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            # Check if next line is not a table line
            if i + 1 >= len(lines) or '|' not in lines[i + 1]:
                # Create table
                if len(table_lines) > 2:  # Header + separator + at least one row
                    # Parse table
                    rows = []
                    for tline in table_lines:
                        if tline.strip().replace('|', '').replace('-', '').replace(' ', ''):
                            cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                            if cells and not all(c.strip('-: ') == '' for c in cells):
                                rows.append(cells)

                    if rows:
                        # Create Word table
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Light Grid Accent 1'

                        for row_idx, row_data in enumerate(rows):
                            for col_idx, cell_data in enumerate(row_data):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data
                                # Bold first row (header)
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True

                        doc.add_paragraph()  # Add space after table
                in_table = False
                table_lines = []

        # Handle lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)

        elif re.match(r'^\d+\.\s', line.strip()):
            # Numbered list
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)

        # Handle regular paragraphs
        elif line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Empty line - add paragraph break
        else:
            if i > 0 and lines[i-1].strip():  # Only if previous line wasn't empty
                doc.add_paragraph()

        i += 1

    # Save document
    doc.save(output_file)
    print(f"Word document created: {output_file}")

if __name__ == "__main__":
    md_file = "/home/user/PanSSR/PanSSR_manuscript.md"
    output_file = "/home/user/PanSSR/PanSSR_manuscript.docx"

    md_content = parse_markdown_file(md_file)
    create_word_document(md_content, output_file)
